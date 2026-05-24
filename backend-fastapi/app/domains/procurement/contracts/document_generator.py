"""
Generación de documento de contrato desde plantilla con sustitución de [VARIABLES].

FASE 6A — Datos completos:
  - Sustituye [VAR] en .docx (python-docx) o .pdf (pdfplumber + reportlab + pypdf)
  - Output siempre PDF
  - Deja contract.status en DRAFT (fase borrador Admin)

Usa generadores ya existentes cuando la plantilla es la que estaba en disco.
"""
from __future__ import annotations

import importlib
import io
import logging
import re
import uuid
from datetime import datetime, timezone
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any, Optional

from fastapi import HTTPException, status
from num2words import num2words
from sqlmodel import Session, select

from app.core.config import settings
from app.platform.contracts_core.models import (
    Contract,
    ContractDocument,
    ContractDocumentType,
    ContractStatus,
    ContractTemplate,
)
from app.domains.documents.storage import build_contract_document_path, ensure_parent_dir
from app.domains.procurement.contracts import crud as contract_crud
from app.domains.procurement.contracts.insurance import (
    compute_subcontract_insurance_amount,
    format_insurance_amount_es,
)

logger = logging.getLogger("app.procurement.document_generator")

_VAR_PATTERN = re.compile(r"\[([A-Z][A-Z0-9_]*)\]")


def _resolve_template_path(template: ContractTemplate, contract: Contract) -> Path:
    """
    Resuelve rutas legacy (Windows host) a la ruta real del runtime (contenedor).
    """
    raw = (template.file_path or "").strip()
    candidate = Path(raw)
    if raw and candidate.exists():
        return candidate

    storage_base = Path(settings.contracts_storage_path)
    normalized = raw.replace("\\", "/")
    marker = "data/contracts/"
    if marker in normalized:
        suffix = normalized.split(marker, 1)[1].lstrip("/")
        mapped = storage_base / Path(suffix)
        if mapped.exists():
            return mapped

    filename = Path(normalized).name
    fallback_candidates = [
        storage_base / f"tenant_{contract.tenant_id}" / "templates" / filename,
        storage_base / "templates" / filename,
    ]
    for fallback in fallback_candidates:
        if fallback.exists():
            return fallback

    return candidate


# ── Formatters ─────────────────────────────────────────────────────────────────

def _format_date_es(value: Any) -> str:
    """Convierte una fecha ISO (yyyy-mm-dd o datetime) a dd/mm/yyyy.

    Si el valor no es parseable se devuelve tal cual (string) o cadena vacía.
    """
    if value is None or value == "":
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y")
    s = str(value).strip()
    if not s:
        return ""
    # Intentar parsear formatos comunes
    for fmt in ("%Y-%m-%d", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%dT%H:%M:%S.%f", "%d/%m/%Y"):
        try:
            return datetime.strptime(s[: len(fmt) + 8], fmt).strftime("%d/%m/%Y")
        except ValueError:
            continue
    return s


def _parse_forma_pago(raw: Any) -> tuple[str, str]:
    """Separa un string de forma de pago en (método, días).

    Acepta entradas tipo "Confirming 60 días", "transferencia 120 días",
    "Otro: <texto libre> 60 días", "CONFIRMING 60", "60 días", etc.
    Devuelve método en minúsculas y días sin sufijo. Cualquiera puede ser "".
    """
    if raw is None:
        return "", ""
    s = str(raw).strip()
    if not s:
        return "", ""
    dias_match = re.search(r"(\d{1,3})\s*d[íi]as?", s, flags=re.IGNORECASE)
    if not dias_match:
        dias_match = re.search(r"\b(\d{1,3})\b", s)
    dias = dias_match.group(1) if dias_match else ""
    metodo_raw = s
    if dias_match:
        metodo_raw = (s[: dias_match.start()] + s[dias_match.end():]).strip()
    metodo_raw = re.sub(r"d[íi]as?", "", metodo_raw, flags=re.IGNORECASE).strip(" :-,.")
    metodo = metodo_raw.lower().strip()
    return metodo, dias


def _int_to_words_es(value: Any) -> str:
    """Entero en letras castellanas (mayúsculas), con apócope.

    Útil para plantillas donde la cifra precede a un sustantivo masculino
    (p.ej. 'UN TRABAJADOR', 'VEINTIÚN TRABAJADORES'). Vacío si no es válido.
    """
    if value is None or value == "":
        return ""
    try:
        n = int(value)
    except (TypeError, ValueError):
        try:
            n = int(Decimal(str(value)))
        except (TypeError, ValueError, InvalidOperation):
            return ""
    try:
        return _apocope_uno(num2words(n, lang="es")).upper()
    except (NotImplementedError, OverflowError, TypeError):
        return ""


_APOCOPE_VEINTIUNO_RE = re.compile(r"\bveintiuno\b")
_APOCOPE_UNO_RE = re.compile(r"\buno\b")


def _apocope_uno(text: str) -> str:
    """Aplica apócope español 'uno' → 'un', 'veintiuno' → 'veintiún'.

    Se aplica:
      - Antes de 'mil' / 'millón' / 'millones' (apócope obligatorio dentro
        de la cifra: 'treinta y un mil', 'veintiún millones').
      - Al final de la cadena, donde concuerda con el sustantivo masculino
        que sigue ('treinta y un euros', 'veintiún trabajadores').
    """
    # 1. Apócope ante 'mil' / 'millón' / 'millones'.
    text = _APOCOPE_VEINTIUNO_RE.sub(
        lambda m: "veintiún" if _followed_by_mil_or_millon(text, m.end()) else m.group(0),
        text,
    )
    text = _APOCOPE_UNO_RE.sub(
        lambda m: "un" if _followed_by_mil_or_millon(text, m.end()) else m.group(0),
        text,
    )
    # 2. Apócope final (precede al sustantivo en la plantilla).
    if text.endswith(" uno"):
        return text[:-4] + " un"
    if text == "uno":
        return "un"
    if text.endswith(" veintiuno"):
        return text[:-10] + " veintiún"
    if text == "veintiuno":
        return "veintiún"
    return text


def _followed_by_mil_or_millon(text: str, pos: int) -> bool:
    rest = text[pos:].lstrip()
    return rest.startswith("mil") or rest.startswith("millón") or rest.startswith("millones")


def _amount_to_words_eur(value: Any) -> str:
    """Importe a letras en euros: 'X EURO(S) CON Y CÉNTIMO(S)' (mayúsculas).

    Solo añade la parte de céntimos si hay decimales no nulos. Aplica
    concordancia singular/plural y apócope (1 EURO, 21 EUROS sin apócope
    en plural pero 21 con apócope cuando va seguido de sustantivo).
    """
    if value is None or value == "":
        return ""
    try:
        amount = Decimal(str(value))
    except (InvalidOperation, TypeError, ValueError):
        return ""
    quantized = amount.quantize(Decimal("0.01"))
    sign = "MENOS " if quantized < 0 else ""
    quantized = abs(quantized)
    euros = int(quantized)
    cents = int((quantized - euros) * 100)
    try:
        euros_text = _apocope_uno(num2words(euros, lang="es")).upper()
        cents_text = _apocope_uno(num2words(cents, lang="es")).upper() if cents else ""
    except (NotImplementedError, OverflowError, TypeError):
        return ""
    euro_unit = "EURO" if euros == 1 else "EUROS"
    result = f"{sign}{euros_text} {euro_unit}"
    if cents:
        cent_unit = "CÉNTIMO" if cents == 1 else "CÉNTIMOS"
        result += f" CON {cents_text} {cent_unit}"
    return result


def _format_milestones(additional: dict, schedule: dict) -> str:
    """Construye la lista de hitos formateada con fechas en dd/mm/yyyy.

    Prioriza `milestones_items` (lista estructurada con name/start/end/description).
    Cae al string plano `milestones`/`hitos` si no hay estructura.
    """
    items = (
        additional.get("milestones_items")
        or additional.get("hitos_items")
        or schedule.get("milestones_items")
    )
    if isinstance(items, list) and items:
        lines: list[str] = []
        for raw in items:
            if not isinstance(raw, dict):
                continue
            name = str(raw.get("name") or raw.get("nombre") or "").strip()
            start = _format_date_es(raw.get("start") or raw.get("inicio"))
            end = _format_date_es(raw.get("end") or raw.get("fin") or raw.get("finalizacion"))
            description = str(
                raw.get("description") or raw.get("descripcion") or raw.get("observaciones") or ""
            ).strip()
            parts: list[str] = []
            if name:
                parts.append(f"{name}:")
            if start:
                parts.append(f"inicio: {start}")
            if end:
                parts.append(f"finalizacion: {end}")
            if description:
                parts.append(f"observaciones: {description}")
            if parts:
                lines.append(" — ".join(parts) if len(parts) > 1 else parts[0])
        if lines:
            return "\n".join(lines)
    # Fallback al texto plano
    plain = (
        additional.get("milestones")
        or additional.get("hitos")
        or additional.get("hitos_clave")
        or schedule.get("milestones")
    )
    return str(plain).strip() if plain else ""


# ── Comparative lines (for HTML SUMINISTRO table) ──────────────────────────────


def _format_eur(value: Any) -> str:
    """Formato '1.234,56 €' (es-ES) para importes del PDF."""
    if value is None or value == "":
        return ""
    try:
        n = float(value)
    except (TypeError, ValueError):
        return ""
    s = f"{n:,.2f}"
    # de "1,234.56" a "1.234,56"
    s = s.replace(",", "X").replace(".", ",").replace("X", ".")
    return f"{s} €"


def _format_qty(value: Any) -> str:
    """Cantidad numérica con coma decimal, sin separador de miles."""
    if value is None or value == "":
        return ""
    try:
        n = float(value)
    except (TypeError, ValueError):
        return ""
    if n == int(n):
        return f"{int(n)}"
    return f"{n:.2f}".replace(".", ",")


def _selected_supplier_name(comparative_data: dict, contract: Contract) -> str:
    """Devuelve el nombre del proveedor ganador (oferta seleccionada).

    Resolución por prioridad: offers[].id == selected_offer_id → supplier_name
    o offer_name. Si no hay match, devuelve "" y el caller hará best-effort.
    """
    sel_id = comparative_data.get("selected_offer_id") or contract.selected_offer_id
    if not sel_id:
        return ""
    offers = comparative_data.get("offers") or []
    for off in offers:
        if not isinstance(off, dict):
            continue
        if off.get("id") == sel_id:
            return str(
                off.get("supplier_name")
                or off.get("offer_name")
                or ""
            ).strip()
    return ""


def _build_suministro_lines(contract: Contract) -> tuple[list[dict], str]:
    """Construye filas y total para la tabla de materiales del SUMINISTRO.

    Cada fila contiene: med, uds, descripcion, precio, importe (todas strings
    ya formateadas en es-ES). Solo se incluyen líneas del proveedor ganador
    con datos económicos (descarta cabeceras sin cantidad/precio).

    Devuelve ([], "") si no hay datos suficientes.
    """
    cd = contract.comparative_data or {}
    lines = cd.get("lines") or []
    if not isinstance(lines, list) or not lines:
        return [], ""

    selected_supplier = _selected_supplier_name(cd, contract)

    def _norm(s: Any) -> str:
        return str(s or "").strip().casefold()

    sel_key = _norm(selected_supplier)

    rows: list[dict] = []
    total = 0.0
    for raw in lines:
        if not isinstance(raw, dict):
            continue
        prices = raw.get("prices") or []
        if not isinstance(prices, list):
            continue
        # Encuentra la entrada de precio del proveedor ganador. Si no hay
        # selected, coge la primera con datos económicos.
        chosen: Optional[dict] = None
        if sel_key:
            for p in prices:
                if isinstance(p, dict) and _norm(p.get("proveedor")) == sel_key:
                    chosen = p
                    break
        if chosen is None and prices:
            for p in prices:
                if isinstance(p, dict) and (
                    p.get("precio_unitario") is not None or p.get("importe") is not None
                ):
                    chosen = p
                    break
        if chosen is None:
            continue

        cantidad = raw.get("cantidad")
        precio_unit = chosen.get("precio_unitario")
        importe = chosen.get("importe")
        # Filtra líneas que no son partidas reales:
        #  - cabeceras del comparativo (sin cantidad, sin precio, sin importe).
        #  - filas de meta/coeficiente del Excel (ej. "% OFERTADO POR
        #    PROVEEDOR") que llevan precio_unitario pero ni cantidad ni
        #    importe; se reconocen por NO tener importe.
        if importe in (None, "") and cantidad in (None, ""):
            continue

        try:
            if importe is not None:
                total += float(importe)
        except (TypeError, ValueError):
            pass

        rows.append({
            "med": _format_qty(cantidad),
            "uds": str(raw.get("unidad") or "").strip(),
            "descripcion": str(raw.get("descripcion") or "").strip(),
            "precio": _format_eur(precio_unit),
            "importe": _format_eur(importe),
        })

    total_str = _format_eur(total) if rows and total > 0 else ""
    return rows, total_str


# ── Context builder ────────────────────────────────────────────────────────────

def _build_substitution_context(
    contract: Contract, session: Optional[Session] = None
) -> dict[str, Any]:
    """
    Builds a flat dict {VAR_NAME: value} covering all spec-required [VARIABLE] names.
    Values are always strings (empty string when missing).

    Si se pasa `session`, consulta la tabla canónica `proveedores` por CIF
    para obtener el representante legal cuando las columnas dedicadas del
    contrato están vacías. Esto neutraliza basura del JSONB legacy.
    """
    contract_data = contract.contract_data or {}
    comparative_data = contract.comparative_data or {}
    economic = contract_data.get("economic") or {}
    schedule = contract_data.get("schedule") or {}
    additional = contract_data.get("additional") or {}
    project_meta = contract_data.get("project") or {}
    header = comparative_data.get("header") if isinstance(comparative_data, dict) else {}
    if not isinstance(header, dict):
        header = {}

    def _str(v: Any) -> str:
        return str(v).strip() if v is not None else ""

    project_name = _str(
        project_meta.get("nombre_obra")
        or header.get("obra_nombre")
        or comparative_data.get("obra")
        or contract.title
    )
    resources_meta = contract_data.get("resources") or {}
    # SUMINISTRO usa columna dedicada `project_number`; priorizarla sobre JSONB.
    project_code = _str(
        contract.project_number
        or project_meta.get("num_obra")
        or resources_meta.get("work_number")
        or header.get("obra_num")
        or contract.project_id
        or contract.id
    )
    worksite_client_name = ""
    if session is not None and project_code:
        try:
            from app.domains.work.catalog_api import get_worksite_by_code_for_tenant

            worksite = get_worksite_by_code_for_tenant(
                session,
                tenant_id=contract.tenant_id,
                code=project_code,
            )
            worksite_client_name = _str(getattr(worksite, "client_name", None))
        except Exception:
            worksite_client_name = ""
    # Forma de pago: SIEMPRE prioridad al comparativo (payment_method_agreed).
    # Está compuesto como "<método> <días> días" o "<otro_texto> <días> días".
    forma_pago_raw = (
        economic.get("payment_method_agreed")
        or additional.get("forma_pago_pactada")
        or economic.get("payment_method")
        or economic.get("forma_pago")
        or additional.get("forma_pago")
    )
    forma_pago_metodo, forma_pago_dias = _parse_forma_pago(forma_pago_raw)
    # SUMINISTRO/SUBCONTRATACIÓN: columnas dedicadas tienen prioridad si están
    # rellenas (las pobla el formulario tras el comparativo).
    if contract.payment_method:
        method_clean = contract.payment_method.strip()
        if method_clean.upper() == "OTROS" and contract.payment_method_other_text:
            forma_pago_metodo = contract.payment_method_other_text.strip().lower()
        elif method_clean:
            forma_pago_metodo = method_clean.lower()
    if contract.payment_days is not None:
        forma_pago_dias = str(contract.payment_days)
    payment_conditions = forma_pago_metodo
    payment_term_from_pactada = (
        f"{forma_pago_dias} días" if forma_pago_dias else ""
    )
    work_description = _str(
        additional.get("units_description")
        or additional.get("descripcion_uds")
        or additional.get("descripcion_unidades")
    )

    # Date parts (use contract creation date or today)
    ref_date = contract.created_at or datetime.now(timezone.utc)
    _MONTHS_ES = [
        "", "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
    ]
    # SUMINISTRO usa columna dedicada `promoter`; priorizarla sobre JSONB.
    promotor = _str(
        worksite_client_name
        or contract.promoter
        or project_meta.get("promotora")
        or project_meta.get("promotor")
        or header.get("promotora")
        or header.get("promotor")
        or contract_data.get("promotora")
        or contract_data.get("promotor")
    )
    plazo = _str(
        schedule.get("duration")
        or schedule.get("plazo")
        or contract_data.get("plazo_ejecucion")
    )

    # Datos del representante/gerente y escritura del proveedor.
    # REGLA DE NEGOCIO: el snapshot del contrato (columnas dedicadas y editables
    # desde el formulario) es la fuente de verdad para ESTE contrato. Las
    # ediciones del form NUNCA tocan la tabla canónica `proveedores`; quedan
    # únicamente como override por contrato.
    # Prioridad:
    #   1. Columnas dedicadas en Contract (override manual del form).
    #   2. Lookup en `proveedores` por CIF (fuente canónica, universal —
    #      el lookup ignora `contract_type` por diseño) — solo como fallback
    #      cuando el snapshot del contrato no tiene dato.
    # No hay fallback al JSONB del comparativo: queda vacío si no hay dato
    # válido en BBDD (mejor mostrar [NOMBRE_GERENTE] vacío que un nombre
    # incorrecto del lado URDECON).
    # Hacemos SIEMPRE lookup en `proveedores` (no solo cuando falta dato) para
    # poder cross-check contra las columnas dedicadas del contrato y detectar
    # drift entre el snapshot guardado y la fuente canónica. La discrepancia
    # se registra como WARNING para auditoría, pero NO altera el valor
    # renderizado — el override del form siempre manda.
    provider_nombre = ""
    provider_nif = ""
    provider_deed_type = ""
    provider_deed_date: Any = None
    provider_notary_name = ""
    provider_notary_protocol = ""
    if session is not None and contract.supplier_tax_id:
        try:
            from app.domains.invoices.ocr.repo import _get_provider_by_tax_id

            provider = _get_provider_by_tax_id(
                session,
                tax_id=contract.supplier_tax_id,
            )
            if provider:
                provider_nombre = _str(provider.get("nombre_gerente"))
                provider_nif = _str(
                    provider.get("nif_gerente") or provider.get("dni_gerente")
                )
                provider_deed_type = _str(provider.get("tipo_escritura"))
                provider_deed_date = provider.get("fecha_escritura")
                provider_notary_name = _str(provider.get("nombre_notario"))
                provider_notary_protocol = _str(provider.get("numero_protocolo"))
        except Exception:
            logger.warning(
                "Lookup canónico `proveedores` falló para contract_id=%s tax_id=%s",
                contract.id,
                contract.supplier_tax_id,
                exc_info=True,
            )

    # Cross-check informativo: si el snapshot del contrato (columnas dedicadas)
    # difiere de `proveedores`, log WARNING para auditoría. NO se altera el
    # valor: el snapshot manda porque representa el override explícito del
    # usuario en el formulario para ESTE contrato.
    contract_nombre = _str(contract.supplier_legal_rep_name)
    contract_nif = _str(contract.supplier_legal_rep_dni)
    if provider_nif and contract_nif and contract_nif != provider_nif:
        logger.warning(
            "Drift gerente contract_id=%s tax_id=%s contract.supplier_legal_rep_dni=%s vs proveedores.nif_gerente=%s — usando snapshot del contrato (override del form)",
            contract.id,
            contract.supplier_tax_id,
            contract_nif,
            provider_nif,
        )
    nombre_gerente = _str(contract_nombre or provider_nombre)
    nif_gerente = _str(contract_nif or provider_nif)

    logistics = contract_data.get("logistics") or {}
    # SUMINISTRO usa columnas dedicadas freight_responsible/unloading_responsible.
    # Legados (additional.freight_party / logistics.shipping_type) como fallback.
    portes = _str(
        contract.freight_responsible
        or additional.get("freight_party")
        or logistics.get("shipping_type")
        or additional.get("portes")
    )
    descargas = _str(
        contract.unloading_responsible
        or additional.get("unloading_party")
        or logistics.get("unloading_type")
        or additional.get("descargas")
    )
    # SUMINISTRO/SUBCONTRATACIÓN tienen columna dedicada milestones_text.
    hitos = _str(contract.milestones_text) or _format_milestones(additional, schedule)
    # SUMINISTRO usa columna dedicada `duration_text`; priorizarla sobre JSONB.
    duracion_obra = _str(
        contract.duration_text
        or project_meta.get("duracion_obra")
        or schedule.get("duration")
        or schedule.get("duration_months")
        or plazo
    )
    termino_pago = _str(
        payment_term_from_pactada
        or additional.get("termino_pago")
        or economic.get("payment_term")
        or economic.get("payment_days")
    )

    # Datos de escritura del proveedor.
    # Cadena: columna dedicada (override manual del form) > lookup canónico
    # en `proveedores` por CIF. NUNCA caer al JSONB del comparativo.
    tipo_escritura = _str(contract.deed_type or provider_deed_type)
    nombre_notario = _str(contract.notary_name or provider_notary_name)
    numero_protocolo = _str(contract.notary_protocol or provider_notary_protocol)
    fecha_escritura = _format_date_es(contract.deed_date or provider_deed_date)

    # Tokens SUBCONTRATACIÓN específicos (precios, trabajadores, garantía).
    # `NUM_TRAB` debe tolerar contratos donde el dato quedó solo en resources.
    raw_num_trab = (
        contract.min_workers
        if contract.min_workers not in (None, "")
        else resources_meta.get("workers_count")
        or resources_meta.get("workers_on_site")
    )
    precio_numero = _str(contract.total_amount)
    precio_letra = _amount_to_words_eur(contract.total_amount)
    num_trab = _str(raw_num_trab)
    num_trab_letra = _int_to_words_es(raw_num_trab)
    seguro = format_insurance_amount_es(
        contract.insurance_amount or compute_subcontract_insurance_amount(contract.total_amount)
    )
    # En SUBCONTRATACION el seguro RC y la garantía son conceptos distintos:
    # no reutilizar la cuantía del seguro como fallback de GARANTIA.
    retention_value = _str(economic.get("retention")).upper()
    if retention_value == "NO":
        retencion_garantia_default = (
            "El CONTRATISTA retendrá el 0% de cada factura que expida el SUBCONTRATISTA, "
            "por tanto, la certificación y/o factura mensual debe llevar reflejada la "
            "retención por garantía."
        )
    elif retention_value == "SI":
        retencion_garantia_default = (
            "El CONTRATISTA practicará una retención del cinco por ciento (5 %) sobre el "
            "importe de cada certificación o factura emitida por el SUBCONTRATISTA, en "
            "concepto de garantía de la correcta ejecución de los trabajos. En "
            "consecuencia, cada certificación y/o factura mensual deberá reflejar "
            "expresamente dicha retención.\n\n"
            "Las cantidades retenidas se mantendrán durante la ejecución de la obra. A la "
            "emisión de la última certificación, el SUBCONTRATISTA deberá aportar un aval "
            "bancario por importe equivalente al cinco por ciento (5 %) del importe final "
            "certificado, con una vigencia mínima de doce (12) meses, que permitirá la "
            "liberación y canje de las retenciones practicadas durante la obra."
        )
    else:
        retencion_garantia_default = ""
    retencion_garantia = _str(contract.warranty_text) or retencion_garantia_default
    garantia = retencion_garantia
    # Token SERVICIOS: tipo de servicio acordado.
    tipo_servicio = _str(contract.service_category)
    # Fecha de fin de obra (alias semántico de FECHA_FIN para plantilla SUBCONTRATACIÓN).
    fecha_fin = _format_date_es(
        schedule.get("end_date") or project_meta.get("fecha_fin") or contract.work_end_date
    )

    ctx: dict[str, str] = {
        # ── Tokens nuevos plantilla SUMINISTRO ───────────────────────────────
        "RAZON_SOCIAL": _str(contract.supplier_name),
        "CIF": _str(contract.supplier_tax_id),
        "DIRECCION_EMPRESA": _str(contract.supplier_address),
        "NOMBRE_GERENTE": nombre_gerente,
        "NIF_GERENTE": nif_gerente,
        "NOMBRE_OBRA": project_name,
        "NUM_OBRA": project_code,
        "NUMERO_OBRA": project_code,
        "PROMOTORA": promotor,
        "DURACION_OBRA": duracion_obra,
        "HITOS": hitos,
        "FORMA_PAGO": payment_conditions,
        "PORTES": portes,
        "DESCARGAS": descargas,
        "TERMINO_PAGO": termino_pago,
        # ── Tokens SUBCONTRATACIÓN ──────────────────────────────────────────
        "TIPO_ESCRITURA": tipo_escritura,
        "FECHA_ESCRITURA": fecha_escritura,
        "NOMBRE_NOTARIO": nombre_notario,
        "NUMERO_PROTOCOLO": numero_protocolo,
        "PRECIO_NUMERO": precio_numero,
        "PRECIO_LETRA": precio_letra,
        "NUM_TRAB": num_trab,
        "NUM_TRAB_LETRA": num_trab_letra,
        "GARANTIA": garantia,
        "RETENCION_GARANTIA": retencion_garantia,
        "SEGURO": seguro,
        "FIN_OBRA": fecha_fin,
        # ── Token SERVICIOS ─────────────────────────────────────────────────
        "TIPO_SERVICIO": tipo_servicio,
        # ── Tokens legados (mantenidos por compatibilidad) ───────────────────
        "NOMBRE_PROVEEDOR": _str(contract.supplier_name),
        "CIF_NIF": _str(contract.supplier_tax_id),
        "DIRECCION": _str(contract.supplier_address),
        "CIUDAD": _str(contract.supplier_city),
        "CODIGO_POSTAL": _str(contract.supplier_postal_code),
        "EMAIL_PROVEEDOR": _str(contract.supplier_email),
        "TELEFONO": _str(contract.supplier_phone),
        "IBAN": _str(contract.supplier_bank_iban),
        "IMPORTE_TOTAL": _str(contract.total_amount),
        "IMPORTE_LETRAS": "",
        "FECHA_INICIO": _format_date_es(
            schedule.get("start_date") or project_meta.get("fecha_inicio") or contract.work_start_date
        ),
        "FECHA_FIN": fecha_fin,
        "CONDICIONES_PAGO": payment_conditions,
        "DESCRIPCION_TRABAJOS": work_description,
        "NOMBRE_PROYECTO": project_name,
        "CODIGO_PROYECTO": project_code,
        "PROMOTOR": promotor,
        "PLAZO_EJECUCION": plazo,
        "DIA": str(ref_date.day),
        "MES": _MONTHS_ES[ref_date.month],
        "ANIO": str(ref_date.year),
        "PAIS": _str(contract.supplier_country),
        "NOMBRE_CONTACTO": _str(contract.supplier_contact_name),
        "BIC": _str(contract.supplier_bank_bic),
        "MONEDA": _str(contract.currency or "EUR"),
    }

    # Líneas/partidas del comparativo: solo se exponen para tipos cuyo PDF
    # (plantilla HTML) las renderiza como tabla. Para los tipos legacy DOCX
    # estas claves se ignoran (Word no consume listas).
    try:
        lineas, total_lineas = _build_suministro_lines(contract)
    except Exception:
        lineas, total_lineas = [], ""
    ctx["LINEAS"] = lineas
    ctx["TOTAL_LINEAS"] = total_lineas

    return ctx


# ── DOCX generation ────────────────────────────────────────────────────────────

_FORCED_FONT_NAME = "Arial"
# Reducción global del tamaño de fuente respecto al de la plantilla.
# Si un run no tiene tamaño explícito, asumimos default 11pt → 9pt resultante.
_FONT_SIZE_DELTA_PT = -2
_FONT_SIZE_DEFAULT_PT = 11
_FONT_SIZE_MIN_PT = 8


def _shrink_font_size(run: Any) -> None:
    """Reduce el tamaño de fuente del run en `_FONT_SIZE_DELTA_PT` puntos.

    - Si el run tiene tamaño explícito, le resta el delta (piso `_FONT_SIZE_MIN_PT`).
    - Si no, asigna `_FONT_SIZE_DEFAULT_PT + _FONT_SIZE_DELTA_PT`.
    """
    try:
        from docx.shared import Pt  # type: ignore

        current = run.font.size
        if current is not None:
            new_pt = max(_FONT_SIZE_MIN_PT, int(current.pt) + _FONT_SIZE_DELTA_PT)
        else:
            new_pt = max(_FONT_SIZE_MIN_PT, _FONT_SIZE_DEFAULT_PT + _FONT_SIZE_DELTA_PT)
        run.font.size = Pt(new_pt)
    except Exception:
        pass


def _force_font_on_run(run: Any) -> None:
    """Aplica la fuente Arial al run y reduce su tamaño en `_FONT_SIZE_DELTA_PT`
    puntos. python-docx por defecto solo setea w:rFonts ascii; para que Word
    respete Arial en todos los rangos de caracteres forzamos también
    east-asia/hAnsi/cs via OxmlElement.
    """
    try:
        run.font.name = _FORCED_FONT_NAME
        rPr = run._element.get_or_add_rPr()
        # Importar perezosamente para no afectar el módulo si python-docx no está.
        from docx.oxml.ns import qn  # type: ignore
        from docx.oxml import OxmlElement  # type: ignore

        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rFonts.set(qn(f"w:{attr}"), _FORCED_FONT_NAME)
    except Exception:
        # Si algo falla, dejar la fuente original; no abortar la generación.
        pass

    _shrink_font_size(run)


def _apply_arial_to_doc(doc: Any) -> None:
    """Aplica Arial al estilo Normal, headers, footers y todas las tablas
    anidadas. Complementa el force-font por run que se hace en _replace_in_paragraph.
    """
    # Estilo Normal: define la fuente y tamaño por defecto del documento
    try:
        from docx.shared import Pt  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.oxml import OxmlElement  # type: ignore

        normal = doc.styles["Normal"]
        normal.font.name = _FORCED_FONT_NAME
        current_size = normal.font.size
        if current_size is not None:
            new_pt = max(_FONT_SIZE_MIN_PT, int(current_size.pt) + _FONT_SIZE_DELTA_PT)
        else:
            new_pt = max(_FONT_SIZE_MIN_PT, _FONT_SIZE_DEFAULT_PT + _FONT_SIZE_DELTA_PT)
        normal.font.size = Pt(new_pt)

        rpr = normal.element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rpr.append(rFonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rFonts.set(qn(f"w:{attr}"), _FORCED_FONT_NAME)
    except Exception:
        pass

    # Headers / footers de todas las secciones
    try:
        for section in doc.sections:
            for container in (section.header, section.footer):
                if container is None:
                    continue
                for para in container.paragraphs:
                    for run in para.runs:
                        _force_font_on_run(run)
                for table in container.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    _force_font_on_run(run)
    except Exception:
        pass


def _replace_in_run(run: Any, replacements: dict[str, str]) -> None:
    text = run.text
    for var, value in replacements.items():
        text = text.replace(f"[{var}]", value)
    run.text = text


def _replace_in_paragraph(para: Any, replacements: dict[str, str]) -> None:
    # Join runs to handle split [VAR] tokens, replace, then rebuild first run
    full_text = "".join(r.text for r in para.runs)
    replaced = full_text
    for var, value in replacements.items():
        replaced = replaced.replace(f"[{var}]", value)
    if replaced != full_text and para.runs:
        para.runs[0].text = replaced
        for run in para.runs[1:]:
            run.text = ""
    # Forzar Arial en todos los runs del párrafo (incluso los que no cambiaron),
    # para uniformar la tipografía del contrato generado.
    for run in para.runs:
        _force_font_on_run(run)


def _generate_from_docx(
    template_path: Path,
    replacements: dict[str, str],
    output_pdf: Path,
) -> bool:
    try:
        docx_module = importlib.import_module("docx")
    except ImportError:
        # Fall back to docxtpl's bundled docx
        try:
            docx_module = importlib.import_module("docxtpl")
            docx_module = importlib.import_module("docx")
        except ImportError:
            logger.error("python-docx not available")
            return False

    try:
        doc = docx_module.Document(str(template_path))
    except Exception as exc:
        logger.error("Cannot open docx template %s: %s", template_path, exc)
        return False

    # Replace in body paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

    # Forzar Arial también en headers/footers y en el estilo Normal por defecto
    _apply_arial_to_doc(doc)

    # Save to temp docx then convert to PDF
    tmp_docx = output_pdf.with_suffix(".tmp.docx")
    try:
        ensure_parent_dir(tmp_docx)
        doc.save(str(tmp_docx))
    except Exception as exc:
        logger.error("Cannot save filled docx: %s", exc)
        return False

    from app.domains.documents.docx_templates import _convert_docx_to_pdf
    ok = _convert_docx_to_pdf(tmp_docx, output_pdf)
    try:
        tmp_docx.unlink(missing_ok=True)
    except Exception:
        pass
    return ok


# ── PDF generation ─────────────────────────────────────────────────────────────

def _generate_from_pdf(
    template_path: Path,
    replacements: dict[str, str],
    output_pdf: Path,
) -> bool:
    """
    Replace [VAR] occurrences in a PDF template using pdfplumber + reportlab + pypdf.
    Each [VAR] token is covered by a white rectangle and the replacement value is
    drawn at the same position.
    """
    try:
        pdfplumber = importlib.import_module("pdfplumber")
        reportlab_canvas = importlib.import_module("reportlab.pdfgen.canvas")
        pypdf_mod = importlib.import_module("pypdf")
    except ImportError as exc:
        logger.error("Missing PDF library: %s", exc)
        return False

    PdfReader = getattr(pypdf_mod, "PdfReader", None)
    PdfWriter = getattr(pypdf_mod, "PdfWriter", None)
    if PdfReader is None or PdfWriter is None:
        return False

    try:
        with pdfplumber.open(str(template_path)) as pdf:
            pages_info = []
            for page in pdf.pages:
                page_height = float(page.height)
                page_width = float(page.width)
                placements: list[tuple[float, float, float, float, str]] = []
                words = page.extract_words(x_tolerance=2, y_tolerance=2) or []

                # Detect [VAR] spanning one or more word tokens
                # Build full-line groups then search
                lines: dict[int, list[dict]] = {}
                for w in words:
                    line_key = int(float(w.get("top", 0)) // 5)
                    lines.setdefault(line_key, []).append(w)

                for line_words in lines.values():
                    line_text = " ".join(str(w.get("text", "")) for w in line_words)
                    for match in _VAR_PATTERN.finditer(line_text):
                        var_name = match.group(1)
                        replacement = replacements.get(var_name, "")
                        if not replacement:
                            continue
                        # Find word(s) that cover this [VAR] match
                        var_token = f"[{var_name}]"
                        matching = [
                            w for w in line_words
                            if var_token in str(w.get("text", ""))
                               or var_name in str(w.get("text", ""))
                        ]
                        if not matching:
                            matching = line_words  # fallback: whole line

                        x0 = min(float(w.get("x0", 0)) for w in matching)
                        x1 = max(float(w.get("x1", x0 + 60)) for w in matching)
                        top = min(float(w.get("top", 0)) for w in matching)
                        bottom = max(float(w.get("bottom", top + 12)) for w in matching)
                        # Convert pdfplumber coords (top-down) to reportlab (bottom-up)
                        rl_y = page_height - bottom
                        height = max(10.0, bottom - top)
                        width = max(60.0, x1 - x0)
                        placements.append((x0, rl_y, width, height, replacement))

                pages_info.append({
                    "width": page_width,
                    "height": page_height,
                    "placements": placements,
                })
    except Exception as exc:
        logger.error("Cannot read PDF template %s: %s", template_path, exc)
        return False

    # Build overlay PDF with replacements
    overlay_buf = io.BytesIO()
    try:
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import pt

        c = rl_canvas.Canvas(overlay_buf)
        for page_info in pages_info:
            c.setPageSize((page_info["width"], page_info["height"]))
            c.setFont("Helvetica", 9)
            for (x0, y, w, h, value) in page_info["placements"]:
                # White cover
                c.setFillColorRGB(1, 1, 1)
                c.rect(x0 - 1, y - 1, w + 2, h + 4, fill=1, stroke=0)
                # Text
                c.setFillColorRGB(0, 0, 0)
                c.drawString(x0, y + 2, value[:120])  # cap length
            c.showPage()
        c.save()
    except Exception as exc:
        logger.error("Cannot build overlay PDF: %s", exc)
        return False

    # Merge overlay onto original
    try:
        overlay_buf.seek(0)
        reader_orig = PdfReader(str(template_path))
        reader_overlay = PdfReader(overlay_buf)
        writer = PdfWriter()

        for i, orig_page in enumerate(reader_orig.pages):
            if i < len(reader_overlay.pages):
                orig_page.merge_page(reader_overlay.pages[i])
            writer.add_page(orig_page)

        ensure_parent_dir(output_pdf)
        with open(str(output_pdf), "wb") as f:
            writer.write(f)
        return True
    except Exception as exc:
        logger.error("Cannot merge overlay PDF: %s", exc)
        return False


# ── Main entry point ───────────────────────────────────────────────────────────

def generate_contract_from_template(
    session: Session,
    *,
    contract: Contract,
    created_by_id: Optional[int],
) -> ContractDocument:
    """
    FASE 6A — Genera el PDF final del contrato sustituyendo [VARIABLES] de la plantilla.
    Deja contract.status en DRAFT (fase borrador editable solo
    por Administración). El paso a PENDING_REVIEW lo dispara Admin con
    `admin_approve_draft` (review_service).
    Crea o actualiza ContractDocument(type=CONTRACT).
    """
    if not contract.template_id:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="El contrato no tiene plantilla asignada.",
        )

    # Asegura que los datos del proveedor estén sincronizados con la tabla
    # Supplier (lookup por CIF) antes de sustituir tokens en la plantilla.
    # Sin esto, contratos creados sin pasar por selección de oferta podrían
    # renderizar [DIRECCION_EMPRESA], [NOMBRE_GERENTE], etc. vacíos.
    try:
        contract_crud.ensure_supplier_snapshot(session, contract=contract)
    except HTTPException:
        raise
    except Exception as exc:
        logger.warning(
            "ensure_supplier_snapshot falló para contract_id=%s: %s",
            contract.id,
            exc,
        )

    template = session.get(ContractTemplate, contract.template_id)
    if not template:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Plantilla del contrato no encontrada.",
        )

    template_path = _resolve_template_path(template, contract)
    if not template_path.exists():
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Archivo de plantilla no encontrado en disco: {template_path}",
        )

    context = _build_substitution_context(contract, session=session)

    output_filename = f"contrato_{contract.id}_{uuid.uuid4().hex[:8]}.pdf"
    output_path = build_contract_document_path(
        contract.tenant_id,
        contract.id,
        ContractDocumentType.CONTRACT,
        output_filename,
    )
    ensure_parent_dir(output_path)

    # Prefer HTML+CSS rendering (Jinja2 + WeasyPrint) when an HTML template
    # exists for this contract type. Avoids DOCX→LibreOffice→PDF reflow issues
    # (misplaced signatures, broken layouts, font drift). Falls back to the
    # legacy DOCX/PDF path for contract types not yet migrated.
    from app.domains.procurement.contracts.html_renderer import (
        HtmlRendererError,
        has_html_template_for,
        render_contract_html_to_pdf,
    )

    if has_html_template_for(contract.type):
        try:
            ok = render_contract_html_to_pdf(
                contract_type=contract.type,
                context=context,
                output_path=output_path,
            )
        except HtmlRendererError as exc:
            logger.error("HTML renderer failed for contract_id=%s: %s", contract.id, exc)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error generando documento (HTML): {exc}",
            )
    else:
        fmt = template.file_format.lower()
        if fmt == "docx":
            ok = _generate_from_docx(template_path, context, output_path)
        elif fmt == "pdf":
            ok = _generate_from_pdf(template_path, context, output_path)
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Formato de plantilla no soportado: {fmt}",
            )

    if not ok:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Error generando documento desde plantilla.",
        )

    # Persist or update ContractDocument
    existing = session.exec(
        select(ContractDocument).where(
            ContractDocument.tenant_id == contract.tenant_id,
            ContractDocument.contract_id == contract.id,
            ContractDocument.doc_type == ContractDocumentType.CONTRACT,
        )
    ).first()

    if existing:
        existing.path = str(output_path)
        existing.created_by_id = created_by_id
        session.add(existing)
        doc = existing
    else:
        doc = ContractDocument(
            tenant_id=contract.tenant_id,
            contract_id=contract.id,
            doc_type=ContractDocumentType.CONTRACT,
            path=str(output_path),
            created_by_id=created_by_id,
        )
        session.add(doc)

    # Tras generar el PDF el contrato se queda en DRAFT
    # (fase borrador editable SOLO por Administración). Cuando Admin pulse
    # "Aprobar" el contrato pasa a PENDING_REVIEW y se abren los slots
    # JURIDICO/JEFE_OBRA/DIRECTOR_TECNICO (el de ADMIN queda ya aprobado).
    if contract.status != ContractStatus.DRAFT:
        contract.status = ContractStatus.DRAFT
    contract.updated_at = datetime.now(timezone.utc)
    session.add(contract)
    contract_crud.ensure_contract_admin_assignment(session, contract=contract)
    session.flush()

    contract_crud._log_event(
        session,
        tenant_id=contract.tenant_id,
        contract_id=contract.id,
        user_id=created_by_id,
        event_type="contract.document_generated",
        payload={"template_id": template.id, "output": str(output_path)},
    )

    session.commit()
    session.refresh(doc)
    session.refresh(contract)
    return doc
